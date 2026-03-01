import json
import logging
import uuid
from io import BytesIO
from typing import AsyncGenerator

from openai import AsyncOpenAI
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import text

from app.config import settings
from app.services.rag import rag_service

log = logging.getLogger(__name__)

EXTRACTION_PROMPT = """You are a legal clause analyzer. Given chunks from a contract, determine if the document contains a clause related to: "{query}".

Return JSON:
{{
  "found": true or false,
  "clause_text": "exact quote from the document (max 500 chars)" or null,
  "summary": "one-sentence plain-language summary of what this clause says" or null,
  "confidence": "high" or "low",
  "chunk_index": which chunk contains the clause (0-based index into the provided chunks)
}}

If no relevant clause is found, return {{"found": false, "clause_text": null, "summary": null, "confidence": "low", "chunk_index": null}}.
{lang_instruction}"""

_LANG_INSTRUCTIONS = {
    "en": "Respond with clause_text and summary in English.",
    "es": "Responde con clause_text y summary en español.",
}


class ClauseComparisonService:
    def __init__(self):
        self._openai = AsyncOpenAI(api_key=settings.openai_api_key)

    async def compare_clauses(
        self,
        db: AsyncSession,
        org_id: uuid.UUID,
        user_id: uuid.UUID,
        query: str,
        project_id: uuid.UUID | None = None,
        language: str = "es",
    ) -> AsyncGenerator[str, None]:
        # 1. Get all visible documents
        scope = "AND d.project_id = CAST(:project_id AS uuid)" if project_id else ""
        visibility = "AND (d.indexed_by_user_id = CAST(:user_id AS uuid) OR d.visibility = 'org')"

        doc_sql = text(f"""
            SELECT d.id, d.title, d.project_id, d.source, d.source_url
            FROM documents d
            WHERE d.org_id = CAST(:org_id AS uuid)
              {scope}
              {visibility}
              AND d.source != 'gmail'
        """)
        params: dict = {"org_id": str(org_id), "user_id": str(user_id)}
        if project_id:
            params["project_id"] = str(project_id)

        result = await db.execute(doc_sql, params)
        all_docs = result.fetchall()
        doc_map = {row.id: row for row in all_docs}

        if not all_docs:
            yield f"data: {json.dumps({'type': 'status', 'total': 0})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            return

        # 2. Embed the clause query
        query_vector = await rag_service._embed_query(query)
        vector_str = "[" + ",".join(str(v) for v in query_vector) + "]"

        # 3. For each document, retrieve the top 5 most relevant chunks
        #    (no score threshold — GPT-4o decides if the clause is present)
        chunks_by_doc: dict[uuid.UUID, list[dict]] = {}

        for doc_row in all_docs:
            doc_id = doc_row.id
            chunk_sql = text("""
                SELECT
                    c.id,
                    c.content,
                    c.chunk_index,
                    1 - (c.embedding <=> CAST(:embedding AS vector)) AS score
                FROM chunks c
                WHERE c.document_id = CAST(:doc_id AS uuid)
                  AND c.embedding IS NOT NULL
                ORDER BY c.embedding <=> CAST(:embedding AS vector)
                LIMIT 5
            """)
            try:
                chunk_result = await db.execute(
                    chunk_sql, {"embedding": vector_str, "doc_id": str(doc_id)}
                )
                rows = chunk_result.fetchall()
                if rows:
                    doc_chunks: list[dict] = []
                    fetched_indices = set()
                    for row in rows:
                        doc_chunks.append({
                            "chunk_id": str(row.id),
                            "content": row.content,
                            "document_id": doc_id,
                            "score": float(row.score),
                            "chunk_index": row.chunk_index,
                            "title": doc_row.title,
                            "source": doc_row.source,
                            "source_url": doc_row.source_url,
                        })
                        fetched_indices.add(row.chunk_index)

                    # Fetch adjacent chunks (chunk_index ± 1) for context
                    adjacent_indices = set()
                    for idx in fetched_indices:
                        adjacent_indices.add(idx - 1)
                        adjacent_indices.add(idx + 1)
                    adjacent_indices -= fetched_indices  # don't re-fetch

                    if adjacent_indices:
                        placeholders = ", ".join(str(int(i)) for i in adjacent_indices if i >= 0)
                        if placeholders:
                            adj_sql = text(f"""
                                SELECT c.id, c.content, c.chunk_index
                                FROM chunks c
                                WHERE c.document_id = CAST(:doc_id AS uuid)
                                  AND c.chunk_index IN ({placeholders})
                            """)
                            adj_result = await db.execute(adj_sql, {"doc_id": str(doc_id)})
                            for adj_row in adj_result.fetchall():
                                doc_chunks.append({
                                    "chunk_id": str(adj_row.id),
                                    "content": adj_row.content,
                                    "document_id": doc_id,
                                    "score": 0.0,
                                    "chunk_index": adj_row.chunk_index,
                                    "title": doc_row.title,
                                    "source": doc_row.source,
                                    "source_url": doc_row.source_url,
                                })

                    chunks_by_doc[doc_id] = doc_chunks
            except Exception as e:
                log.warning("Chunk retrieval failed for doc %s: %s", doc_id, e)

        total = len(all_docs)
        yield f"data: {json.dumps({'type': 'status', 'total': total})}\n\n"

        # 4. Process each document via GPT-4o
        lang_instruction = _LANG_INSTRUCTIONS.get(language, _LANG_INSTRUCTIONS["en"])
        prompt = EXTRACTION_PROMPT.format(query=query, lang_instruction=lang_instruction)
        current = 0

        for doc_id, doc_chunks in chunks_by_doc.items():
            current += 1
            doc_info = doc_map.get(doc_id)
            if not doc_info:
                continue

            # Sort by chunk_index for coherent reading order, but keep top-scored first
            sorted_chunks = sorted(doc_chunks, key=lambda c: c.get("score", 0), reverse=True)[:10]
            doc_text = "\n\n".join(
                f"[CHUNK {i}]\n{c['content']}" for i, c in enumerate(sorted_chunks)
            )

            try:
                response = await self._openai.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": prompt},
                        {"role": "user", "content": doc_text},
                    ],
                    temperature=0.1,
                    max_tokens=800,
                    response_format={"type": "json_object"},
                )
                raw = response.choices[0].message.content.strip()
                parsed = json.loads(raw)

                chunk_idx = parsed.get("chunk_index")
                matched_chunk_id = None
                if chunk_idx is not None and 0 <= chunk_idx < len(sorted_chunks):
                    matched_chunk_id = sorted_chunks[chunk_idx]["chunk_id"]

                yield f"data: {json.dumps({'type': 'result', 'data': {'document_id': str(doc_id), 'title': doc_info.title, 'project_id': str(doc_info.project_id) if doc_info.project_id else None, 'found': parsed.get('found', False), 'clause_text': parsed.get('clause_text'), 'summary': parsed.get('summary'), 'confidence': parsed.get('confidence', 'low'), 'chunk_id': matched_chunk_id, 'source': doc_info.source, 'source_url': doc_info.source_url}})}\n\n"
            except Exception as e:
                log.error("Clause extraction failed for doc %s: %s", doc_id, e)
                yield f"data: {json.dumps({'type': 'result', 'data': {'document_id': str(doc_id), 'title': doc_info.title, 'project_id': str(doc_info.project_id) if doc_info.project_id else None, 'found': False, 'clause_text': None, 'summary': None, 'confidence': 'low', 'chunk_id': None, 'source': doc_info.source, 'source_url': doc_info.source_url}})}\n\n"

        # 5. Emit documents that had no chunks retrieved
        for doc_id in doc_map:
            if doc_id not in chunks_by_doc:
                doc_info = doc_map[doc_id]
                yield f"data: {json.dumps({'type': 'missing', 'document_id': str(doc_id), 'title': doc_info.title})}\n\n"

        # 6. Done
        yield f"data: {json.dumps({'type': 'done'})}\n\n"

    def generate_docx(self, query: str, results: list[dict], missing: list[dict], language: str = "es") -> BytesIO:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title_text = "Comparación de cláusulas" if language == "es" else "Clause Comparison"
        title = doc.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        query_label = "Cláusula buscada" if language == "es" else "Clause searched"
        doc.add_paragraph(f"{query_label}: {query}").italic = True

        # Summary table
        found_results = [r for r in results if r.get("found")]
        header_doc = "Documento" if language == "es" else "Document"
        header_summary = "Resumen" if language == "es" else "Summary"
        header_confidence = "Confianza" if language == "es" else "Confidence"

        if found_results:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = header_doc
            hdr[1].text = header_summary
            hdr[2].text = header_confidence

            for r in found_results:
                row = table.add_row().cells
                row[0].text = r.get("title", "")
                row[1].text = r.get("summary", "") or ""
                row[2].text = r.get("confidence", "low")

        # Per-document detail
        detail_heading = "Detalle por documento" if language == "es" else "Detail by document"
        doc.add_heading(detail_heading, level=2)

        for r in found_results:
            doc.add_heading(r.get("title", ""), level=3)
            if r.get("clause_text"):
                p = doc.add_paragraph()
                p.add_run(r["clause_text"]).font.color.rgb = RGBColor(0x57, 0x53, 0x4E)

        # Missing section
        if missing:
            missing_heading = "Cláusula no encontrada" if language == "es" else "Clause not found"
            doc.add_heading(missing_heading, level=2)
            not_found_results = [r for r in results if not r.get("found")]
            all_missing = missing + not_found_results
            for m in all_missing:
                doc.add_paragraph(f"• {m.get('title', '')}")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def generate_pdf(self, query: str, results: list[dict], missing: list[dict], language: str = "es") -> BytesIO:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Use built-in font with latin-1 encoding for broad compatibility
        pdf.set_font("Helvetica", "B", 16)
        title_text = "Comparacion de clausulas" if language == "es" else "Clause Comparison"
        pdf.cell(0, 12, title_text, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        pdf.set_font("Helvetica", "I", 11)
        query_label = "Clausula buscada" if language == "es" else "Clause searched"
        pdf.cell(0, 8, f"{query_label}: {query}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(6)

        found_results = [r for r in results if r.get("found")]

        # Summary table
        if found_results:
            pdf.set_font("Helvetica", "B", 10)
            col_widths = [70, 80, 30]
            headers = (
                ["Documento", "Resumen", "Confianza"]
                if language == "es"
                else ["Document", "Summary", "Confidence"]
            )
            for i, h in enumerate(headers):
                pdf.cell(col_widths[i], 8, h, border=1)
            pdf.ln()

            pdf.set_font("Helvetica", "", 9)
            for r in found_results:
                x_start = pdf.get_x()
                y_start = pdf.get_y()
                pdf.cell(col_widths[0], 8, (r.get("title", "") or "")[:40], border=1)
                pdf.cell(col_widths[1], 8, (r.get("summary", "") or "")[:50], border=1)
                pdf.cell(col_widths[2], 8, r.get("confidence", "low"), border=1)
                pdf.ln()
            pdf.ln(6)

        # Detail
        pdf.set_font("Helvetica", "B", 13)
        detail_heading = "Detalle por documento" if language == "es" else "Detail by document"
        pdf.cell(0, 10, detail_heading, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for r in found_results:
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 8, r.get("title", ""), new_x="LMARGIN", new_y="NEXT")
            if r.get("clause_text"):
                pdf.set_font("Helvetica", "", 9)
                pdf.multi_cell(0, 5, r["clause_text"])
            pdf.ln(4)

        # Missing
        not_found_results = [r for r in results if not r.get("found")]
        all_missing = missing + not_found_results
        if all_missing:
            pdf.set_font("Helvetica", "B", 13)
            missing_heading = "Clausula no encontrada" if language == "es" else "Clause not found"
            pdf.cell(0, 10, missing_heading, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            for m in all_missing:
                pdf.cell(0, 7, f"- {m.get('title', '')}", new_x="LMARGIN", new_y="NEXT")

        buf = BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return buf


clause_comparison_service = ClauseComparisonService()
